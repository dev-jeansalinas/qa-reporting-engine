import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def cargar_casos_prueba(ruta):
    try:
        with open(ruta, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError) as e:
        print(f"[-] Error en datos: {e}")
        return []


def generar_word_snva(datos_hu, info_proyecto, ruta_plantilla, ruta_salida):
    # Cargar la plantilla corporativa
    doc = Document(ruta_plantilla)

    # Normalizar datos: Si es una lista de casos plana, envolverla en una estructura de HU
    if datos_hu and "pruebas" not in datos_hu[0]:
        datos_hu = [
            {
                "hu_id": info_proyecto.get("modulo", "N/A"),
                "hu_nombre": "General",
                "pruebas": datos_hu,
            }
        ]

    # Preparar cálculos para la tabla de resumen
    total_casos = sum(len(hu.get("pruebas", [])) for hu in datos_hu)
    exitosos = sum(
        sum(1 for c in hu.get("pruebas", []) if c.get("estado") == "Exitoso")
        for hu in datos_hu
    )
    fallidos = sum(
        sum(1 for c in hu.get("pruebas", []) if c.get("estado") == "Fallido")
        for hu in datos_hu
    )
    solucionados = sum(
        sum(1 for c in hu.get("pruebas", []) if c.get("estado") == "Solucionado")
        for hu in datos_hu
    )

    # Mapa de etiquetas { } vs Datos del JSON
    reemplazos = {
        "{CLIENTE}": info_proyecto.get("cliente", "N/A"),
        "{VERSION}": str(info_proyecto.get("version", "N/A")),
        "{ALOJAMIENTO}": str(info_proyecto.get("alojamiento", "N/A")),
        "{ANALISTA}": info_proyecto.get("tester", "N/A"),
        "{FECHA_INICIO}": info_proyecto.get("inicio_pruebas", "N/A"),
        "{FECHA_FIN}": info_proyecto.get("fin_pruebas", "N/A"),
        "{TOTAL}": str(total_casos),
        "{EXITOSOS}": str(exitosos),
        "{FALLIDOS}": str(fallidos),
        "{SOLUCIONADOS}": str(solucionados),
    }

    # Reemplazo automático en todas las tablas del documento
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for etiqueta, valor in reemplazos.items():
                    if etiqueta in celda.text:
                        celda.text = celda.text.replace(etiqueta, valor)

    # Generar el detalle de evidencias (Una página por HU)
    titulo_seccion = doc.add_heading("Evidencias de pruebas", level=1)
    titulo_seccion.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in titulo_seccion.runs:
        run.font.name = "Arial"
        run.bold = True
        run.font.size = Pt(11)

    for hu in datos_hu:
        id_hu = hu.get("hu_id", info_proyecto.get("modulo", "N/A"))
        nombre_hu = hu.get("hu_nombre", "N/A")
        id_display = f"{id_hu} - {nombre_hu}"

        for i, caso in enumerate(hu.get("pruebas", []), start=1):
            id_caso = f"TC-{i:02d}"
            nombre_caso = caso.get("nombre_caso", "N/A")
            id_automatizado = f"{id_caso} {nombre_caso}"

            titulo_caso = doc.add_heading(f"{id_automatizado}", level=3)
            for run in titulo_caso.runs:
                run.font.name = "Arial"
                run.bold = True
                run.font.size = Pt(11)
            
            espacio_imagen = doc.add_paragraph("") 
            run_espacio = espacio_imagen.add_run("")
            run_espacio.font.name = "Arial"
            run_espacio.font.size = Pt(11)

            resultado_texto = f"Resultado obtenido: {caso.get('resultado_obtenido', '')}"
            parrafo_res = doc.add_paragraph(resultado_texto)
            for run in parrafo_res.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)


        doc.add_page_break()

    doc.save(ruta_salida)
